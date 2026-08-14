"""Text extraction per MIME (§10). Native-speed libraries are optional
extras — every extractor degrades gracefully and records what ran.

Core (always available): text/*, json/yaml/xml, csv/tsv.
Optional: PDF via pypdfium2, DOCX via python-docx, HTML via selectolax
(regex-strip fallback), XLSX via python-calamine.
"""

import asyncio
import csv
import io
import re
from pathlib import Path

import structlog

log = structlog.get_logger("retinue.rag.extract")

MAX_TEXT_CHARS = 4_000_000  # ~1M tokens; enough for any honest RAG corpus

_TAG_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.DOTALL | re.IGNORECASE)
_ANY_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"[ \t]+\n")


class ExtractionError(Exception):
    pass


def _decode(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_html(data: bytes) -> tuple[str, str]:
    try:
        from selectolax.parser import HTMLParser

        tree = HTMLParser(_decode(data))
        for node in tree.css("script,style,noscript"):
            node.decompose()
        return tree.body.text(separator="\n") if tree.body else tree.text(), "selectolax"
    except ImportError:
        stripped = _TAG_RE.sub(" ", _decode(data))
        return _ANY_TAG_RE.sub(" ", stripped), "regex-strip"


def _extract_pdf(path: Path) -> tuple[str, str]:
    try:
        import pypdfium2 as pdfium  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ExtractionError("PDF extraction requires the [extract] extra (pypdfium2)") from exc
    doc = pdfium.PdfDocument(str(path))
    try:
        pages: list[str] = []
        for i in range(len(doc)):
            page = doc[i]
            textpage = page.get_textpage()
            try:
                pages.append(f"[page {i + 1}]\n{textpage.get_text_bounded()}")
            finally:
                textpage.close()
                page.close()
        return "\n\n".join(pages), "pypdfium2"
    finally:
        doc.close()


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("DOCX extraction requires the [extract] extra (python-docx)") from exc
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts), "python-docx"


def _extract_xlsx(path: Path) -> tuple[str, str]:
    try:
        from python_calamine import CalamineWorkbook
    except ImportError as exc:
        raise ExtractionError(
            "XLSX extraction requires the [extract] extra (python-calamine)"
        ) from exc
    workbook = CalamineWorkbook.from_path(str(path))
    sheets: list[str] = []
    for name in workbook.sheet_names:
        rows = workbook.get_sheet_by_name(name).to_python()
        body = "\n".join("\t".join(str(cell) for cell in row) for row in rows[:10_000])
        sheets.append(f"[sheet {name}]\n{body}")
    return "\n\n".join(sheets), "python-calamine"


def _extract_csv(data: bytes, delimiter: str) -> tuple[str, str]:
    reader = csv.reader(io.StringIO(_decode(data)), delimiter=delimiter)
    lines = ["\t".join(row) for _, row in zip(range(50_000), reader, strict=False)]
    return "\n".join(lines), "csv"


def extract_text_sync(path: Path, mime: str, filename: str) -> tuple[str, str]:
    """Returns (text, extractor_name). Raises ExtractionError when unsupported."""
    if mime == "application/pdf":
        text, extractor = _extract_pdf(path)
    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text, extractor = _extract_docx(path)
    elif mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        text, extractor = _extract_xlsx(path)
    elif mime == "text/html":
        text, extractor = _extract_html(path.read_bytes())
    elif mime in ("text/csv", "text/tab-separated-values"):
        delimiter = "\t" if mime.endswith("values") else ","
        text, extractor = _extract_csv(path.read_bytes(), delimiter)
    elif mime.startswith("text/") or mime in (
        "application/json",
        "application/xml",
        "application/yaml",
        "application/x-yaml",
    ):
        text, extractor = _decode(path.read_bytes()), "text"
    elif mime.startswith("image/"):
        raise ExtractionError("image files carry no extractable text (OCR not configured)")
    else:
        raise ExtractionError(f"no extractor for {mime}")

    text = _WS_RE.sub("\n", text).strip()
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
    return text, extractor


async def extract_text(path: Path, mime: str, filename: str) -> tuple[str, str]:
    return await asyncio.to_thread(extract_text_sync, path, mime, filename)
